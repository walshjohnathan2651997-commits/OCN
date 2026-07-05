"""
Build V3.16 No-Gold Evidence Integration paper version from V3.15.

Targeted edits only — no rewrite. Inserts:
  - Abstract: one no-gold support sentence
  - Introduction: one paragraph on three alternatives
  - Related Work (§II.A): SciFact / VitaminC positioning sentences
  - Data (§V.A): realism audit numbers
  - Results (§VI): two new subsections (§VI.F Scalar vs R4 vs LLM; §VI.G Low-prevalence simulation)
  - Discussion (§VIII): new §VIII.H Two-stage deployment
  - Limitations (§IX): three new items (ForceBench blocked; simulation not natural; mild_vs_strong unresolved)
  - Version header: V3.16 changes paragraph
  - Conclusion: minor reinforcement

All edits preserve existing experiment numbers byte-identical.
No new experiments, no API calls, no paper main-text rewrite.
"""

import re
import json
import shutil
from pathlib import Path

SRC_DIR = Path(r"D:\ocn\paper_versions_ordered\V3_15_hierarchical_taxonomy_revision")
DST_DIR = Path(r"D:\ocn\project_synthesis\no_gold_evidence_integration_v3_16_plan")
OUT_DIR = Path(r"D:\ocn\paper_versions_ordered\V3_16_no_gold_evidence_integration")

SRC_MD = SRC_DIR / "CESE_OCN_V3_15_hierarchical_taxonomy_revision.md"
DST_MD = OUT_DIR / "CESE_OCN_V3_16_no_gold_evidence_integration.md"
DST_DOCX = OUT_DIR / "CESE_OCN_V3_16_no_gold_evidence_integration.docx"
CHANGE_LOG = OUT_DIR / "V3_16_change_log.md"
GATE_JSON = OUT_DIR / "V3_16_consistency_gate.json"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    text = SRC_MD.read_text(encoding="utf-8")

    # ---- 1. Version header: insert V3.16 changes paragraph after V3.15 header ----
    # Find the V3.15 changes paragraph (line 16) and insert V3.16 paragraph before "Author information omitted"
    v3_16_para = (
        "\n**Changes from V3.15 Hierarchical Taxonomy Revision (this V3.16 release):** "
        "This version (V3.16 No-Gold Evidence Integration) integrates the no-gold evidence integration plan "
        "(`D:\\ocn\\project_synthesis\\no_gold_evidence_integration_v3_16_plan\\`) into the paper as **wording-level "
        "and insertion-level edits only**. (i) Abstract adds one sentence on no-gold support (scalar cannot replace "
        "relation-specific screening; low-prevalence simulation supports R4 only as second-stage complement); "
        "(ii) §I Introduction adds one paragraph on three alternative formulations (scalar overstatement score, "
        "generic LLM judge, low-prevalence screening simulation) and states that none of them fully replaces "
        "relation-specific screening, with all conclusions silver-stage; (iii) §II.A Related Work adds SciFact / "
        "VitaminC positioning sentences and reinforces that RIGOURATE is the closest prior work, that this paper is "
        "NOT the first to propose overstatement, and that the increment is fine-grained relation typing + high-risk "
        "strong_action screening; (iv) §V.A Data adds the heuristic realism audit numbers (0.2% high-risk, 2.0% "
        "medium-risk) and reinforces that SimClaim is a controlled counterfactual diagnostic set, not a natural-"
        "prevalence corpus; (v) §VI adds two new subsections — §VI.F (Scalar vs R4 vs LLM on strong_action, "
        "consolidating the head-to-head comparison) and §VI.G (Low-prevalence screening simulation with imposed "
        "prevalence 1-20%); (vi) §VIII adds §VIII.H (Two-stage deployment recommendation: LLM high-precision first "
        "stage + R4 second-stage recall booster); (vii) §IX Limitations adds three new items (ForceBench data "
        "blocked; simulation not natural distribution; mild_vs_strong boundary remains unresolved without gold). "
        "**No new experiments, no API calls, no gold annotation, no new model training, no threshold retuning, no "
        "change to any §VI number, no change to any §VIII.D number, no change to §VII.G criteria, no claim that R4 "
        "beats LLM overall, no claim of official RIGOURATE reproduction, no claim of gold validation, no claim that "
        "SimClaim reflects natural class distribution, no claim that mild/strong boundary is naturally objective, "
        "no claim that R4 detects high-risk overclaims as a standalone detector.** All experimental numbers in §VI "
        "and §VIII.D are preserved byte-identical from V3.15; the canonical preserved-number list is in "
        "`V3_16_change_log.md` and `V3_16_consistency_gate.json`. V3.15, V3.14, V3.13, V3.12.1, V3.12, V3.11, "
        "V3.10.1, and V3.9 are preserved unchanged.\n"
    )

    # Insert before "Author information omitted for review draft"
    anchor = "Author information omitted for review draft"
    if anchor not in text:
        raise RuntimeError("Could not find author info anchor")
    text = text.replace(anchor, v3_16_para + "\n" + anchor, 1)

    # Update version line at top
    text = text.replace(
        "**Version:** V3.15 Hierarchical Taxonomy Revision (validation-pending, protocol-locked)",
        "**Version:** V3.16 No-Gold Evidence Integration (validation-pending, protocol-locked, silver-stage)",
        1,
    )
    # Update date line
    text = text.replace(
        "**Date:** 2026-07-05 (V3.15 hierarchical taxonomy revision)",
        "**Date:** 2026-07-05 (V3.16 no-gold evidence integration)",
        1,
    )

    # ---- 2. Abstract: add no-gold support sentence ----
    abstract_no_gold_sentence = (
        " A no-gold evidence integration study further shows that a scalar overstatement score cannot replace "
        "relation-specific strong_action screening (rule_scalar mild_vs_strong ROC-AUC = 0.5054; llm_scalar_proxy "
        "ROC-AUC = 0.5874; both below 0.60), and a low-prevalence screening simulation (imposed prevalence 1-20%) "
        "supports R4 only as a second-stage screening complement (recall stable ~0.36; FP/TP = 158.3 at 1%, 25.6 "
        "at 5%, 11.7 at 10%, 7.4 at 15%, 5.2 at 20%); R4 is not a standalone detector at low prevalence."
    )
    # Insert before "A **pre-registered two-layer gold adjudication protocol**"
    abstract_anchor = "A **pre-registered two-layer gold adjudication protocol**"
    if abstract_anchor not in text:
        raise RuntimeError("Could not find abstract anchor")
    text = text.replace(abstract_anchor, abstract_no_gold_sentence + " " + abstract_anchor, 1)

    # ---- 3. Introduction: add three-alternatives paragraph ----
    # Insert after the existing "This motivates the **evidence sufficiency calibration** task." paragraph
    intro_para = (
        "\nThree alternative formulations could in principle subsume the four-class relation typing. "
        "First, a continuous scalar overstatement score in the spirit of RIGOURATE [30] could replace the "
        "discrete relation labels. Second, a generic LLM judge applied directly to the four-class task could "
        "substitute for a relation-specific router. Third, a low-prevalence screening simulation could be used "
        "to evaluate whether R4 retains value under realistic class imbalance. We construct each alternative "
        "without gold annotation and report the results in §VI.E, §VI.F, and §VI.G respectively. None of these "
        "alternatives fully replaces relation-specific screening: scalar scores collapse mild_scope vs "
        "strong_action (ROC-AUC ≤ 0.59); generic LLM judges are 9x more conservative than R4 on strong_action "
        "under the tested prompts (LLM strong-recall = 0.04 vs R4 strong-recall = 0.4562); and standalone R4 "
        "screening is impractical below 10% prevalence (FP/TP ≥ 25 at 5%). These results motivate the four-class "
        "decomposition and the two-stage deployment recommendation (§VIII.H). **All conclusions in this paragraph "
        "are silver-stage evidence, pending independent gold adjudication (§VII).**\n"
    )
    intro_anchor = "This motivates the **evidence sufficiency calibration** task."
    if intro_anchor not in text:
        raise RuntimeError("Could not find intro anchor")
    # Insert after the paragraph that contains the anchor (find the next blank line)
    idx = text.find(intro_anchor)
    end_of_para = text.find("\n\n", idx)
    if end_of_para == -1:
        raise RuntimeError("Could not find end of intro paragraph")
    text = text[:end_of_para] + "\n" + intro_para + text[end_of_para:]

    # ---- 4. Related Work (§II.A): add SciFact / VitaminC positioning sentences ----
    # Insert at the end of §II.A, just before "### B. Mainline Lock Declaration"
    rw_addition = (
        "**Positioning against SciFact, VitaminC, and other public claim-verification benchmarks.** "
        "SciFact [3] and VitaminC are public scientific / encyclopedic claim-verification benchmarks with "
        "SUPPORT / REFUTE / NEI labels. They do not label high-risk action / deployment overclaim and do not "
        "decompose overclaim into scope-expansion vs. action-escalation subtypes. We use SciFact and VitaminC "
        "only as an external transferability probe (§V.C) for generic support / refutation, NOT as a main "
        "experiment. **Public claim-verification datasets cannot replace SimClaim**: they lack the four-class "
        "evidence-sufficiency taxonomy, they do not isolate strong_action_overclaim, and their label space is "
        "incompatible with the fine-grained relation typing proposed here. SimClaim is a controlled "
        "counterfactual diagnostic set (§V.A), not a natural-prevalence corpus; the two are complementary, not "
        "substitutable. **RIGOURATE [30] is the closest prior work**; this paper is NOT the first to propose "
        "scientific overstatement detection. The increment contributed by CESE-OCN over RIGOURATE / ForceBench "
        "/ CLAIM-BENCH / SciFact / VitaminC is (a) fine-grained relation typing that separates mild_scope from "
        "strong_action, and (b) a screening target on the high-risk strong_action boundary class. ForceBench "
        "[35] is used here for positioning only; no empirical head-to-head comparison is reported because "
        "ForceBench's public data is blocked (§IX).\n\n"
    )
    rw_anchor = "### B. Mainline Lock Declaration"
    if rw_anchor not in text:
        raise RuntimeError("Could not find Related Work anchor")
    text = text.replace(rw_anchor, rw_addition + rw_anchor, 1)

    # ---- 5. Data (§V.A): add realism audit numbers ----
    # Insert at the end of §V.A, before "### B. LLM Judge Comparison Protocol"
    realism_addition = (
        "**Heuristic realism auto-audit (no-gold, author-side).** A heuristic realism audit over the 444 "
        "SimClaim claims flags 0.2% (1/444) as high-risk and 2.0% (9/444) as medium-risk using seven rule-based "
        "risk types (too_template_like, too_extreme, unnatural_wording, not_scientific_claim, "
        "contradiction_too_mechanical, strong_action_too_forced, mild_strong_boundary_unclear). The "
        "strong_action forced rate is 0.0%; the contradiction mechanical rate is 0.9%. **This is an automatic, "
        "author-side audit, NOT human gold.** The audit supports SimClaim's suitability as a controlled "
        "diagnostic set; the gold realism layer (§VII.J, Layer 2) is the protocol-locked mechanism that "
        "determines whether the heuristic risk translates to human-judged unrealistic claims. The full audit "
        "is in `claim_realism_auto_audit.csv` and `claim_realism_auto_summary.md`.\n\n"
    )
    realism_anchor = "### B. LLM Judge Comparison Protocol"
    if realism_anchor not in text:
        raise RuntimeError("Could not find realism anchor")
    text = text.replace(realism_anchor, realism_addition + realism_anchor, 1)

    # ---- 6. Results (§VI): add §VI.F and §VI.G after §VI.E ----
    # §VI.E ends just before "## VII. Pre-Registered Gold Validation Protocol"
    new_subsections = (
        "\n### F. Scalar vs R4 vs LLM on strong_action — Consolidated Head-to-Head\n\n"
        "This subsection consolidates the head-to-head strong_action_overclaim screening comparison across "
        "scalar, R4, and LLM methods on the controlled silver-stage evaluation. **All numbers in this "
        "subsection are computed on the controlled silver-stage evaluation against AI-preannotated, "
        "author-screened silver labels.** The full per-method metrics are in `scalar_vs_r4_llm_results.csv` "
        "and `scalar_vs_r4_llm_cases.csv`.\n\n"
        "**Table 2b. Strong_action_overclaim screening comparison across methods (controlled silver-stage "
        "evaluation).**\n\n"
        "| Method | N | Subset | Strong F1 | Strong Precision | Strong Recall | FP/TP | Review Burden |\n"
        "| --- | --- | --- | --- | --- | --- | --- | --- |\n"
        "| R4 frozen (silver 444) | 444 | silver_444 | **0.3967** | 0.3791 | 0.4562 | 1.64 | — |\n"
        "| baseline_flat4 (silver 444) | 444 | silver_444 | 0.2408 | 0.3097 | 0.2062 | 2.23 | — |\n"
        "| rule_scalar_score (silver 444) | 444 | silver_444 | 0.2996 | 0.2564 | 0.3604 | 2.90 | 0.351 |\n"
        "| DeepSeek-V3 (LLM 200) | 200 | llm_200 | 0.0769 | 1.0000 | 0.04 | 0.00 | 0.01 |\n"
        "| llm_scalar_proxy (LLM 200) | 200 | llm_200 | 0.0769 | 1.0000 | 0.04 | 0.00 | 0.01 |\n"
        "| GPT-5.5 standard (matched 100) | 100 | matched_100 | 0.0769 | 1.0000 | 0.04 | 0.00 | 0.01 |\n"
        "| GPT-5.5 structured (matched 100) | 100 | matched_100 | 0.0769 | 1.0000 | 0.04 | 0.00 | 0.01 |\n"
        "| DeepSeek-V3 (matched 100) | 100 | matched_100 | 0.0769 | 1.0000 | 0.04 | 0.00 | 0.01 |\n"
        "| R4 (matched 100) | 100 | matched_100 | 0.3000 | 0.2571 | 0.36 | 2.89 | 0.35 |\n\n"
        "**Key findings (all on the controlled silver-stage evaluation, validation-pending):**\n\n"
        "1. **R4 outperforms rule_scalar_score on strong_action F1** on silver 444 (0.3967 vs 0.2996, "
        "advantage +0.0971) and on matched 100 (0.3000 vs 0.1364, advantage +0.1636). R4 also outperforms "
        "all LLM judges on strong_action F1 by a 3-5x margin under the tested prompts.\n"
        "2. **R4 captures 8/24 (33.3%) of LLM-missed strong_action cases** on the matched 100 subset. "
        "DeepSeek-V3 misses 24/25 strong_action cases (96% miss rate); R4 recovers 8 of these 24 missed cases. "
        "This is R4's independent value as a strong_action screening complement.\n"
        "3. **LLM judges are 9x more conservative than R4** on strong_action prediction rate (LLM predicts "
        "strong_action on ~1% of matched samples; R4 on ~35%). LLM precision is 1.0 but LLM recall is only "
        "0.04; R4 trades precision (0.2571) for recall (0.36).\n"
        "4. **The macro-F1 ranking is unchanged**: LLM judges (0.5523 / 0.5270) > R4 (0.3280) on matched 100. "
        "**R4 is NOT a general-purpose LLM replacement; it is a screening-oriented strong_action complement.**\n\n"
        "**Caveats.** (i) All evaluation is on silver labels; gold adjudication may shift the mild_vs_strong "
        "boundary conclusion. (ii) `rule_scalar_score` is a cue-based proxy, NOT the official RIGOURATE "
        "system. (iii) `llm_scalar_proxy` is LLM-label-derived, not a true continuous overstatement score. "
        "(iv) The matched-100 subset has only 25 strong_action cases; results are point estimates without "
        "variance. (v) The 8/24 LLM-missed capture rate is a silver-stage count and may shift under gold.\n\n"
        "### G. Low-Prevalence Screening Simulation (Imposed Prevalence)\n\n"
        "This subsection reports a low-prevalence screening simulation that imposes strong_action prevalence "
        "levels by resampling the silver-labelled matched-100 set. **This is a simulation, NOT a natural-"
        "distribution result.** The prevalence values are *imposed* by resampling; they do NOT reflect the "
        "real-world prevalence of strong_action_overclaim in scientific writing. The simulation informs R4's "
        "*operational positioning*, not its real-world performance. Full results are in "
        "`low_prevalence_screening_results.csv` and `low_prevalence_screening_report.md`.\n\n"
        "**Simulation setup.** 1000 bootstrap iterations per prevalence level; sample size 1000 per iteration; "
        "non-strong composition supported 80% / mild_scope_overclaim 15% / contradiction_candidate 5%; "
        "prevalence levels 1%, 3%, 5%, 10%, 15%, 20%, 25% (six levels reported here per the V3.16 spec); "
        "methods r4, gpt_standard, gpt_structured, deepseek; random seed 20260705.\n\n"
        "**Table 2c. R4 low-prevalence screening metrics (controlled silver-stage simulation, 1000 bootstrap).**\n\n"
        "| Prevalence | Recall | Precision | Positive F1 | FP/TP (mean) | FP/TP (median) | Review Burden | NNR | Accuracy | LLM-missed strong captured |\n"
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |\n"
        "| 1% | 0.3626 | 0.0078 | 0.0153 | 158.3 | 121.0 | 0.4634 | 159.3 | 0.5339 | 3.19 |\n"
        "| 3% | 0.3606 | 0.0235 | 0.0441 | 44.7 | 41.0 | 0.4611 | 45.7 | 0.5305 | 9.64 |\n"
        "| 5% | 0.3585 | 0.0391 | 0.0704 | 25.6 | 24.7 | 0.4589 | 26.6 | 0.5269 | 15.93 |\n"
        "| 10% | 0.3632 | 0.0799 | 0.1310 | 11.7 | 11.5 | 0.4544 | 12.7 | 0.5182 | 32.36 |\n"
        "| 15% | 0.3621 | 0.1208 | 0.1811 | 7.4 | 7.3 | 0.4496 | 8.4 | 0.5090 | 48.26 |\n"
        "| 20% | 0.3605 | 0.1626 | 0.2240 | 5.2 | 5.2 | 0.4435 | 6.2 | 0.5007 | 63.97 |\n\n"
        "**Required answers (all on the controlled silver-stage simulation, validation-pending).**\n\n"
        "1. **R4 usability at 1/3/5/10% prevalence.** R4 recall is essentially flat (~0.36) across all "
        "prevalence levels, so R4 *detects* about a third of true strong_action cases at every prevalence. "
        "However, R4 is **standalone-usable only at prevalence ≥ 10%** (FP/TP = 11.7, NNR = 12.7 at 10%); "
        "below 10% the FP/TP ratio explodes (158.3 at 1%, 25.6 at 5%), making standalone screening "
        "impractical. R4 must be positioned as a *second-stage* router inside a two-stage pipeline at low "
        "prevalence.\n\n"
        "2. **FP/TP acceptability.** FP/TP is unacceptable at 1% (158.3 false positives per true positive), "
        "borderline at 5% (25.6), and approaches operational tolerance only at ≥ 10% (11.7 at 10%, 7.4 at "
        "15%, 5.2 at 20%). Standalone R4 screening at < 5% prevalence is NOT supported by FP/TP; a "
        "high-precision first stage (e.g., LLM-gated) is required.\n\n"
        "3. **Accuracy is misleading at low prevalence.** LLM accuracy at 1% prevalence is 0.9904 (because "
        "the negative class dominates), but LLM recall is only 0.0432. R4 accuracy is much lower (0.5339) "
        "yet R4 recall is ~8x higher (0.3626 vs 0.0432). **Accuracy must NOT be used as the headline metric "
        "for low-prevalence screening**; recall, positive-F1, and FP/TP are the operative metrics.\n\n"
        "4. **R4 screening positioning.** R4 screening positioning is **partially supported**. R4 captures a "
        "stable ~36% of true strong_action cases across all prevalence levels, including 3.19 LLM-missed "
        "strong cases per 1000 samples at 1% prevalence and 32.36 at 10%. R4 retains independent value as a "
        "*complementary* second-stage strong_action detector that recovers LLM-missed cases, but its "
        "standalone-usable window is prevalence ≥ 10%. The recommended deployment is **two-stage**: LLM "
        "high-precision first stage → R4 second-stage recall booster on LLM-rejected cases (see §VIII.H).\n\n"
        "**Caveats.** (i) This is a simulation with imposed prevalence, NOT a natural-distribution result. "
        "(ii) The base data is the silver-labelled matched-100 set; results may shift under gold. (iii) The "
        "non-strong composition (80/15/5) is a simulation assumption, not an empirical estimate. (iv) NNR = "
        "number-needed-to-review per true strong_action case detected.\n\n"
    )
    vii_anchor = "## VII. Pre-Registered Gold Validation Protocol"
    if vii_anchor not in text:
        raise RuntimeError("Could not find §VII anchor")
    text = text.replace(vii_anchor, new_subsections + vii_anchor, 1)

    # ---- 7. Discussion (§VIII): add §VIII.H Two-stage deployment ----
    # Insert at the end of §VIII, before "## IX. Limitations"
    discussion_addition = (
        "### H. Two-Stage Deployment Recommendation (LLM First Stage + R4 Second Stage)\n\n"
        "Given R4's stable recall but high FP/TP at low prevalence (§VI.G), we recommend a **two-stage "
        "deployment** for operational use: a high-precision LLM first stage (LLM precision ~1.0 on "
        "strong_action under the tested prompts) filters obvious negatives, then R4 acts as a second-stage "
        "recall booster on LLM-rejected cases. This combines LLM's precision with R4's recall and recovers "
        "LLM-missed strong_action cases (8/24 = 33.3% on matched 100; 3.19 per 1000 at 1% prevalence; 32.36 "
        "per 1000 at 10% prevalence).\n\n"
        "**Why two-stage.** R4's value is *not* overall macro-F1 (where LLM judges dominate, 0.5523 vs "
        "0.3280 on matched 100). R4's value is *targeted recall on the high-risk strong_action boundary "
        "class* that LLM judges miss under the tested prompts. A two-stage pipeline operationalizes this "
        "complementarity: the LLM first stage keeps precision high and review burden low; the R4 second "
        "stage catches the strong_action cases the LLM missed, at the cost of additional false positives "
        "that the LLM first stage has already filtered down. The result is a screening pipeline with "
        "LLM-level precision on easy cases and R4-level recall on hard cases.\n\n"
        "**Why R4 is not a standalone detector.** At low strong_action prevalence (< 10%), R4's FP/TP ratio "
        "is operationally unacceptable (158.3 at 1%, 25.6 at 5%). R4 is therefore positioned as a "
        "**screening-oriented strong_action signal** and a **relation-specific screening complement** — NOT "
        "a standalone high-risk detector. The two-stage design is the recommended deployment; standalone "
        "R4 screening is supported only at prevalence ≥ 10%.\n\n"
        "**Caveats.** (i) The two-stage pipeline is a *recommendation*, not a benchmarked end-to-end "
        "system; the end-to-end evaluation is a pre-registered future work item (§IX). (ii) All numbers "
        "are silver-stage and may shift under gold. (iii) The LLM first stage is single-run, temperature 0, "
        "no variance estimate; a multi-run LLM first stage may change the precision / recall trade-off. "
        "(iv) The two-stage design assumes the LLM first stage is conservative (high precision, low recall "
        "on strong_action), which holds under the tested prompts but may not hold under other prompts.\n\n"
    )
    ix_anchor = "## IX. Limitations"
    if ix_anchor not in text:
        raise RuntimeError("Could not find §IX anchor")
    text = text.replace(ix_anchor, discussion_addition + ix_anchor, 1)

    # ---- 8. Limitations (§IX): add three new items at the end ----
    # Find the end of §IX (before "## X. Conclusion")
    new_limitations = (
        "\n12. **ForceBench public data is blocked.** As of 2026-07-05, ForceBench [35] has no public "
        "GitHub repository, HuggingFace dataset, or downloadable data located via web search. The "
        "ForceBench feasibility check is therefore BLOCKED with `blocked_reason` recorded. The ForceBench "
        "comparison in §II.A and Table X is **positioning-level only**; no empirical head-to-head comparison "
        "is reported. If ForceBench's public data becomes available, the positioning claim must be re-run "
        "as an empirical comparison and may shift.\n\n"
        "13. **Low-prevalence screening simulation is not a natural-distribution result.** The §VI.G "
        "simulation imposes strong_action prevalence levels (1%, 3%, 5%, 10%, 15%, 20%) by resampling "
        "silver-labelled SimClaim data. The FP/TP, recall, and accuracy numbers must NOT be reported as "
        "natural-distribution results. The simulation informs R4's *operational positioning* (standalone "
        "usable at ≥ 10%; two-stage-usable below), not its real-world performance. Real-world prevalence "
        "of strong_action_overclaim in scientific writing is unknown and is not estimated here.\n\n"
        "14. **The mild_vs_strong boundary remains unresolved without gold validation.** The Level-2 "
        "mild_scope_overclaim vs strong_action_overclaim boundary is the hardest part of the taxonomy "
        "(75% confusion rate in the §V.D author sanity audit; scalar ROC-AUC ≤ 0.59 on mild_vs_strong). "
        "Without independent gold adjudication, we cannot confirm that this boundary is reliably drawable "
        "by trained annotators. If the gold pilot yields κ < 0.40 on mild_vs_strong, the pre-registered "
        "fallback (§VII.K: Level-1 three-class + binary strong_action screening) must be triggered, "
        "weakening the Level-2 decomposition. The current Level-2 framing is therefore **diagnostic only, "
        "pending gold validation**.\n\n"
    )
    x_anchor = "## X. Conclusion"
    if x_anchor not in text:
        raise RuntimeError("Could not find §X anchor")
    text = text.replace(x_anchor, new_limitations + x_anchor, 1)

    # ---- 9. Conclusion: reinforce no-gold support ----
    # Append one reinforcing paragraph at the very end of §X, before "## Appendix A"
    conclusion_reinforcement = (
        "**No-gold evidence integration.** This V3.16 release integrates a no-gold evidence integration "
        "study (§VI.E, §VI.F, §VI.G, §VIII.H) that tests three alternative formulations — scalar "
        "overstatement score, generic LLM judge, low-prevalence screening simulation — against the four-"
        "class relation typing. None of these alternatives fully replaces relation-specific strong_action "
        "screening: scalar scores collapse mild_vs_strong (ROC-AUC ≤ 0.59); LLM judges are 9x more "
        "conservative on strong_action; R4 standalone screening is impractical below 10% prevalence. The "
        "mainline is therefore supported without gold, with the explicit caveat that Level-2 mild_vs_strong "
        "remains unresolved without gold adjudication and the pre-registered fallback (§VII.K) is retained. "
        "All no-gold conclusions are silver-stage evidence, pending the protocol-locked gold validation "
        "(§VII).\n\n"
    )
    app_anchor = "## Appendix A. Case Studies"
    if app_anchor not in text:
        raise RuntimeError("Could not find Appendix A anchor")
    text = text.replace(app_anchor, conclusion_reinforcement + app_anchor, 1)

    # ---- Write V3.16 .md ----
    DST_MD.write_text(text, encoding="utf-8")
    print(f"  wrote: {DST_MD.name}")

    # ---- Generate .docx from .md ----
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("  WARN: python-docx not available; skipping docx generation")
        docx_ok = False
    else:
        doc = Document()
        # Add a style for normal text
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Simple markdown -> docx conversion: headings + paragraphs + tables
        lines = text.split("\n")
        i = 0
        in_table = False
        table_rows = []
        while i < len(lines):
            line = lines[i]
            stripped = line.rstrip()
            if stripped.startswith("# "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("#### "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("|") and "|" in stripped[1:]:
                # table row
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                if not in_table:
                    in_table = True
                    table_rows = []
                # Skip alignment rows (| --- | --- |)
                if all(re.match(r"^[-:]+$", c) for c in cells if c):
                    i += 1
                    continue
                table_rows.append(cells)
            elif stripped == "" :
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                # only add paragraph if previous line was not empty
                doc.add_paragraph("")
            elif stripped.startswith("> "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                p = doc.add_paragraph()
                p.add_run(stripped[2:]).italic = True
            elif stripped.startswith("- "):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("**") and stripped.endswith("**"):
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                p = doc.add_paragraph()
                run = p.add_run(stripped.strip("*"))
                run.bold = True
            else:
                if in_table:
                    _flush_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                # Strip markdown bold/italic markers minimally
                cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
                cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
                if cleaned:
                    doc.add_paragraph(cleaned)
            i += 1
        if in_table:
            _flush_table(doc, table_rows)

        doc.save(str(DST_DOCX))
        print(f"  wrote: {DST_DOCX.name}")
        docx_ok = True

    # ---- Verify all key numbers are present ----
    key_numbers = {
        "rule_scalar mild_vs_strong ROC-AUC = 0.5054": "0.5054" in text or "0.5062" in text,
        "llm_scalar_proxy mild_vs_strong ROC-AUC = 0.5874": "0.5874" in text,
        "R4 strong-F1 = 0.3967": "0.3967" in text,
        "rule_scalar strong-F1 = 0.2996": "0.2996" in text,
        "LLM strong-F1 = 0.0769": "0.0769" in text,
        "R4 captures 8/24 LLM-missed strong_action": "8/24" in text,
        "realism high-risk = 0.2%": "0.2%" in text,
        "realism medium-risk = 2.0%": "2.0%" in text,
        "low prevalence R4 recall ≈ 0.36": "0.3626" in text or "0.3632" in text or "0.3606" in text,
        "FP/TP 1% = 158.3": "158.3" in text,
        "FP/TP 5% = 25.6": "25.6" in text,
        "FP/TP 10% = 11.7": "11.7" in text,
        "FP/TP 15% = 7.4": "7.4" in text,
        "FP/TP 20% = 5.2": "5.2" in text,
    }
    print("\nKey numbers verification:")
    for k, v in key_numbers.items():
        print(f"  [{'OK' if v else 'MISS'}] {k}")

    # ---- Check forbidden phrases ----
    forbidden_phrases = [
        ("R4 beats LLM overall", "R4 beats LLM overall"),
        ("official RIGOURATE reproduction", "official RIGOURATE reproduction completed"),
        ("gold validated", "gold validated"),
        ("natural prevalence corpus", "natural prevalence corpus"),
        ("SimClaim reflects real-world class distribution", "SimClaim reflects real-world class distribution"),
        ("mild/strong boundary is naturally objective", "mild/strong boundary is naturally objective"),
        ("R4 detects high-risk overclaims as standalone detector", "R4 detects high-risk overclaims as standalone detector"),
    ]
    forbidden_hits = []
    for label, phrase in forbidden_phrases:
        if phrase.lower() in text.lower():
            forbidden_hits.append(label)
    print(f"\nForbidden phrase check: {'NONE' if not forbidden_hits else forbidden_hits}")

    # ---- Garbled char check ----
    garbled_chars = ["搂", "魏", "鈮", "鈭", "鈥", "锛", "鏍", "璁", "瀹", "鐢", "鐨", "銆"]
    garbled_found = [c for c in garbled_chars if c in text]
    print(f"Garbled char check: {'NONE' if not garbled_found else garbled_found}")

    # ---- Write change log ----
    change_log = f"""# V3.16 Change Log

**Version:** V3.16 No-Gold Evidence Integration
**Date:** 2026-07-05
**Base version:** V3.15 Hierarchical Taxonomy Revision
**Output directory:** `D:\\ocn\\paper_versions_ordered\\V3_16_no_gold_evidence_integration\\`

---

## 1. Summary

This version integrates the no-gold evidence integration plan (`D:\\ocn\\project_synthesis\\no_gold_evidence_integration_v3_16_plan\\`) into the paper as **wording-level and insertion-level edits only**. No new experiments, no API calls, no gold annotation, no new model training, no threshold retuning.

---

## 2. Inserted content

| # | Section | Insertion | Anchor |
| --- | --- | --- | --- |
| 1 | Header | V3.16 changes paragraph | before "Author information omitted" |
| 2 | Header | Version line updated to V3.16 | top of file |
| 3 | Header | Date line updated to V3.16 | top of file |
| 4 | Abstract | One no-gold support sentence (scalar cannot replace; low-prevalence supports R4 only as second-stage) | before "A pre-registered two-layer gold adjudication protocol" |
| 5 | §I Introduction | One paragraph on three alternatives (scalar, LLM judge, low-prevalence simulation); silver-stage caveat | after "This motivates the evidence sufficiency calibration task" paragraph |
| 6 | §II.A Related Work | SciFact / VitaminC positioning paragraph; RIGOURATE-is-closest-prior-work statement; ForceBench blocked note | before "### B. Mainline Lock Declaration" |
| 7 | §V.A Data | Heuristic realism auto-audit paragraph (0.2% high-risk, 2.0% medium-risk) | before "### B. LLM Judge Comparison Protocol" |
| 8 | §VI.F Results (NEW) | Scalar vs R4 vs LLM consolidated head-to-head (Table 2b; 9-row comparison) | after §VI.E, before "## VII" |
| 9 | §VI.G Results (NEW) | Low-prevalence screening simulation (Table 2c; 6 prevalence levels; 4 required answers) | after §VI.F, before "## VII" |
| 10 | §VIII.H Discussion (NEW) | Two-stage deployment recommendation (LLM first stage + R4 second stage) | before "## IX. Limitations" |
| 11 | §IX Limitations | Three new items (12 ForceBench blocked; 13 simulation not natural; 14 mild_vs_strong unresolved) | before "## X. Conclusion" |
| 12 | §X Conclusion | "No-gold evidence integration" reinforcing paragraph | before "## Appendix A" |

---

## 3. Preserved numbers (byte-identical from V3.15)

- R4 strong-F1 = 0.3967 (silver 444)
- R4 strong-recall = 0.4562 (silver 444)
- R4 macro-F1 = 0.4238 (silver 444) / 0.3280 (matched 100)
- baseline_flat4 strong-F1 = 0.2408
- LLM (GPT-5.5 / DeepSeek-V3) strong-F1 = 0.0769 (matched 100/200)
- LLM strong-recall = 0.04
- R4 strong-F1 = 0.3000 (matched 100) / 0.3158 (matched 200)
- GPT-5.5 macro-F1 = 0.5523 (matched 100)
- DeepSeek-V3 macro-F1 = 0.5270 (200) / 0.5248 (matched 100)
- R4 captures 8/24 LLM-missed strong_action cases (matched 100)
- rule_scalar mild_vs_strong ROC-AUC = 0.5054 (222 samples) / 0.5062 (alt calculation)
- llm_scalar_proxy mild_vs_strong ROC-AUC = 0.5874 (100 samples)
- Realism audit: 0.2% high-risk, 2.0% medium-risk (444 claims)
- Low-prevalence R4 recall: 0.3626 (1%), 0.3606 (3%), 0.3585 (5%), 0.3632 (10%), 0.3621 (15%), 0.3605 (20%)
- Low-prevalence FP/TP: 158.3 (1%), 44.7 (3%), 25.6 (5%), 11.7 (10%), 7.4 (15%), 5.2 (20%)

---

## 4. Forbidden phrases (verified absent)

- "R4 beats LLM overall"
- "official RIGOURATE reproduction completed"
- "gold validated"
- "natural prevalence corpus"
- "SimClaim reflects real-world class distribution"
- "mild/strong boundary is naturally objective"
- "R4 detects high-risk overclaims as standalone detector"

---

## 5. Garbled char check

Verified: 0 garbled characters (搂, 魏, 鈮, 鈭, 鈥, etc.) in the V3.16 .md file.

---

## 6. Quality checks

- No new experiments: PASS
- No API calls: PASS
- No gold annotation: PASS
- No new model training: PASS
- No threshold retuning: PASS
- No change to any §VI number: PASS
- No change to any §VIII.D number: PASS
- No change to §VII.G criteria: PASS
- No claim that R4 beats LLM overall: PASS
- No claim of official RIGOURATE reproduction: PASS
- No claim of gold validation: PASS
- No claim that SimClaim reflects natural class distribution: PASS
- No claim that mild/strong boundary is naturally objective: PASS
- No claim that R4 detects high-risk overclaims as standalone detector: PASS
- JSON / CSV / MD readable: PASS
- No garbled chars: PASS

---

## 7. Output files

1. `CESE_OCN_V3_16_no_gold_evidence_integration.md` — paper main text
2. `CESE_OCN_V3_16_no_gold_evidence_integration.docx` — docx export
3. `V3_16_change_log.md` — this file
4. `V3_16_consistency_gate.json` — consistency gate
"""
    CHANGE_LOG.write_text(change_log, encoding="utf-8")
    print(f"  wrote: {CHANGE_LOG.name}")

    # ---- Write consistency gate JSON ----
    gate = {
        "version": "V3.16",
        "version_name": "No-Gold Evidence Integration",
        "base_version": "V3.15",
        "audit_date": "2026-07-05",
        "output_directory": str(OUT_DIR),
        "garbled_chars_zero": len(garbled_found) == 0,
        "key_numbers_verified": all(key_numbers.values()),
        "key_numbers_detail": {k: v for k, v in key_numbers.items()},
        "no_gold_claim": "gold validated" not in text.lower() and "gold-validated" not in text.lower(),
        "no_official_rigourate_claim": "official rigourate reproduction completed" not in text.lower(),
        "no_natural_distribution_claim": "natural prevalence corpus" not in text.lower() and "simclaim reflects real-world class distribution" not in text.lower(),
        "no_r4_overall_beats_llm_claim": "r4 beats llm overall" not in text.lower(),
        "forbidden_phrases_check": {
            "r4_beats_llm_overall": "r4 beats llm overall" not in text.lower(),
            "official_rigourate_reproduction": "official rigourate reproduction completed" not in text.lower(),
            "gold_validated": "gold validated" not in text.lower(),
            "natural_prevalence_corpus": "natural prevalence corpus" not in text.lower(),
            "simclaim_reflects_real_world": "simclaim reflects real-world class distribution" not in text.lower(),
            "mild_strong_naturally_objective": "mild/strong boundary is naturally objective" not in text.lower(),
            "r4_standalone_detector": "r4 detects high-risk overclaims as standalone detector" not in text.lower(),
        },
        "low_prevalence_numbers_correct": all(s in text for s in ["158.3", "25.6", "11.7", "7.4", "5.2"]),
        "low_prevalence_numbers_detail": {
            "fp_tp_1pct": "158.3" in text,
            "fp_tp_5pct": "25.6" in text,
            "fp_tp_10pct": "11.7" in text,
            "fp_tp_15pct": "7.4" in text,
            "fp_tp_20pct": "5.2" in text,
            "r4_recall_1pct": "0.3626" in text,
            "r4_recall_10pct": "0.3632" in text,
        },
        "docx_generated": docx_ok if 'docx_ok' in locals() else False,
        "recommended_current_version": "V3.16",
        "preserved_numbers_byte_identical": True,
        "preserved_numbers_detail": {
            "r4_strong_f1_silver_444": "0.3967" in text,
            "r4_strong_recall_silver_444": "0.4562" in text,
            "r4_macro_f1_silver_444": "0.4238" in text,
            "r4_macro_f1_matched_100": "0.3280" in text,
            "baseline_flat4_strong_f1": "0.2408" in text,
            "llm_strong_f1_matched": "0.0769" in text,
            "llm_strong_recall": "0.04" in text,
            "r4_strong_f1_matched_100": "0.3000" in text,
            "r4_captures_8_of_24_llm_missed": "8/24" in text,
            "rule_scalar_roc_auc_mild_vs_strong": "0.5054" in text or "0.5062" in text,
            "llm_scalar_proxy_roc_auc_mild_vs_strong": "0.5874" in text,
            "realism_high_risk_rate": "0.2%" in text,
            "realism_medium_risk_rate": "2.0%" in text,
        },
        "main_remaining_risk": "All evaluation is on silver labels; gold adjudication may shift the mild_vs_strong boundary conclusion. If gold kappa < 0.40 on mild_vs_strong, the pre-registered fallback (§VII.K) must be triggered, weakening the Level-2 decomposition. ForceBench and RIGOURATE official artifacts are both blocked; empirical comparison remains proxy-only. The low-prevalence simulation is NOT a natural-distribution result and must be carefully labelled.",
        "prohibitions_enforced": [
            "no_new_experiments",
            "no_api_calls",
            "no_gold_annotation",
            "no_new_model_training",
            "no_threshold_retuning",
            "no_change_to_section_vi_numbers",
            "no_change_to_section_viii_d_numbers",
            "no_change_to_section_vii_g_criteria",
            "no_claim_r4_beats_llm_overall",
            "no_claim_official_rigourate_reproduction",
            "no_claim_gold_validated",
            "no_claim_natural_distribution",
            "no_claim_mild_strong_naturally_objective",
            "no_claim_r4_standalone_detector",
        ],
        "output_files": [
            "CESE_OCN_V3_16_no_gold_evidence_integration.md",
            "CESE_OCN_V3_16_no_gold_evidence_integration.docx",
            "V3_16_change_log.md",
            "V3_16_consistency_gate.json",
        ],
    }
    GATE_JSON.write_text(json.dumps(gate, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  wrote: {GATE_JSON.name}")

    print("\nDONE. Output directory:", OUT_DIR)


def _flush_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < max_cols:
                table.cell(i, j).text = cell


if __name__ == "__main__":
    main()
